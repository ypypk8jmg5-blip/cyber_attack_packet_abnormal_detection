from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "feedback_driven_adaptive_mlops_nad_draft.docx"

ACCENT = "1F4E79"
LIGHT = "EAF2F8"
MID = "D6EAF8"
DARK = RGBColor(31, 78, 121)
MUTED = RGBColor(96, 96, 96)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips=9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_bottom_border(paragraph, color="B7B7B7", size="6") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Feedback-Driven Adaptive MLOps for Network Anomaly Detection"
    hp.style = doc.styles["Header"]
    hp.runs[0].font.size = Pt(9)
    hp.runs[0].font.color.rgb = MUTED
    set_bottom_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Draft manuscript")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in [
        ("Title", 22, 0, 10),
        ("Subtitle", 11, 0, 12),
        ("Heading 1", 16, 14, 6),
        ("Heading 2", 13, 10, 4),
        ("Heading 3", 11, 8, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if "Heading" in name or name == "Title":
            style.font.bold = True
            style.font.color.rgb = DARK


def add_para(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, sum(widths))
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = DARK

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.size = Pt(9)

    doc.add_paragraph()


def build_doc() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    setup_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("A Feedback-Driven Adaptive MLOps Framework for Flow-Level Network Anomaly Detection")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Manuscript Draft | Prepared from the anomaly-detection-mlops prototype | May 2026")

    meta = [
        ["Document type", "Initial research manuscript draft"],
        ["Scope", "Synthetic flow-feature network anomaly detection with adaptive generation and multi-agent fusion"],
        ["Status", "Draft for technical review, experiment expansion, and citation completion"],
    ]
    add_table(doc, ["Field", "Description"], meta, [2200, 7160])

    doc.add_heading("Abstract", level=1)
    add_para(doc,
             "Network anomaly detection systems often degrade when models are trained on static datasets that do not reflect evolving attack variants, imbalanced class distributions, or high-load benign traffic. This paper presents a feedback-driven adaptive MLOps framework for flow-level network anomaly detection. The proposed framework integrates synthetic packet-flow generation, automated model training, per-attack evaluation, adaptive sample reweighting, deterministic signature overrides, and a multi-agent decision-fusion pipeline. Instead of treating dataset generation as a one-time preprocessing step, the framework uses evaluation feedback to increase the representation and difficulty of attack classes with weak recall. A prototype implementation was evaluated on synthetic flow-feature data containing five benign traffic classes and fourteen anomaly classes. Preliminary experiments show that the closed-loop pipeline can reach target performance in a small number of cycles, with a representative single-cycle run achieving F1 = 0.9702, recall = 0.9607, and precision = 0.9799. In the real-time simulation phase, the detector processed one 50-flow batch and produced eight alerts, matching the number of detected anomalies. The results suggest that adaptive generation and operational feedback can improve robustness in synthetic flow-level network anomaly detection, while also highlighting the need for validation on real network telemetry.")
    add_para(doc, "Keywords: network anomaly detection; MLOps; adaptive data generation; synthetic traffic; multi-agent detection; intrusion detection; flow-level security analytics.")

    doc.add_heading("1. Introduction", level=1)
    add_para(doc,
             "Machine-learning-based network anomaly detection has become an attractive approach for identifying suspicious traffic patterns, including denial-of-service attacks, scanning, brute force attempts, command-and-control communication, and data exfiltration. However, model performance is often evaluated under static assumptions: a fixed dataset, a fixed train-test split, and a fixed attack distribution. In operational environments, these assumptions are fragile. Attack patterns shift, benign traffic becomes more diverse, and rare attack classes may remain underrepresented.")
    add_para(doc,
             "This work investigates a practical alternative: an adaptive MLOps loop that continuously connects generation, training, evaluation, and detection. The prototype does not claim to be a production packet capture engine. Rather, it is a flow-feature-based research framework designed to test whether evaluation feedback can guide subsequent synthetic data generation and improve class-specific detection behavior.")
    add_para(doc, "The main contributions are:")
    add_bullets(doc, [
        "A feedback-driven synthetic flow generation method that increases the sampling weight of attack classes with low recall.",
        "A boundary-sample strategy that creates difficult attack variants near normal traffic ranges.",
        "A closed-loop MLOps pipeline that automatically trains, evaluates, promotes the best model, and transitions into a simulated real-time detection phase.",
        "A multi-agent decision architecture combining machine-learning inference, rule signatures, behavioral evidence, threshold management, and alert generation.",
        "A preliminary experimental analysis over fourteen synthetic anomaly classes and five benign traffic classes."
    ])

    doc.add_heading("2. Background and Motivation", level=1)
    add_para(doc,
             "Flow-level anomaly detection summarizes network behavior through aggregate features such as duration, packet rate, byte rate, connection count, failed attempts, destination-port diversity, and outbound ratio. This representation is computationally lighter than raw packet inspection and aligns with common telemetry sources such as NetFlow, Zeek logs, firewall records, and proxy logs. However, the usefulness of flow-level detection depends heavily on the quality and coverage of the training distribution.")
    add_para(doc,
             "In static synthetic datasets, models may learn highly separable attack signatures. For example, a botnet command-and-control class generated only on unusual ports may be easy to detect, yet fail when the same behavior is moved to HTTPS. Similarly, data exfiltration generated only as high-volume transfer may not capture slow-drip exfiltration. These problems motivate an adaptive generator that responds to measured weaknesses, rather than repeatedly producing the same distribution.")

    doc.add_heading("3. Proposed Framework", level=1)
    add_para(doc,
             "The proposed framework consists of two main phases. Phase 1 performs adaptive learning: packet-flow data generation, model training, evaluation, and feedback. Phase 2 performs simulated real-time detection: incoming batches are processed by the trained detector and alerts are emitted for anomalies. The architecture is designed to keep training metrics, per-attack recall, model artifacts, session summaries, and alert summaries synchronized.")

    add_table(doc,
              ["Component", "Role", "Current implementation"],
              [
                  ["Packet-flow generator", "Creates labeled benign and anomalous flow-feature records.", "generate_packets.py and AdaptivePacketGenerator"],
                  ["Trainer", "Fits the primary ML classifier and stores a model bundle.", "RandomForest with scaler and feature list"],
                  ["Evaluator", "Computes global metrics and per-attack recall.", "F1, recall, precision, accuracy, log loss"],
                  ["Adaptive feedback", "Uses weak recall to adjust future attack sampling.", "Recall-target gap weighting"],
                  ["Live detector", "Loads best_model.pkl and classifies incoming flow batches.", "detect_anomaly.py"],
                  ["Signature override", "Forces high-confidence detection for deterministic C2 port evidence.", "C2 ports: 4444, 6667, 1080, 8443, 9001"],
                  ["Multi-agent layer", "Fuses multiple analysis votes before final thresholding.", "32-agent pipeline prototype"]
              ],
              [1900, 3600, 3860])

    doc.add_heading("3.1 Adaptive Packet-Flow Generation", level=2)
    add_para(doc,
             "The adaptive generator reads feedback from recent evaluation artifacts and maintains a per-attack weight vector. If the recall of an attack class falls below the target recall, its sampling weight is increased. The current weighting rule is:")
    add_para(doc, "weight = 1.0 + (target_recall - observed_recall) x boost_factor")
    add_para(doc,
             "where target_recall is 0.90 and boost_factor is 3.0 in the current prototype. The weights are normalized so that the total attack budget remains stable while weak classes receive a larger share of the next training cycle.")
    add_para(doc,
             "The generator also modifies the difficulty of weak attack classes. For example, port scans may use fewer distinct destination ports, HTTP floods may use lower packet rates, DNS tunneling may use packet sizes closer to normal DNS traffic, and botnet C2 may include normal-looking ports such as 80 and 443. This makes the generated samples less trivially separable.")

    doc.add_heading("3.2 Boundary Sampling", level=2)
    add_para(doc,
             "In addition to the standard benign and attack classes, the adaptive generator injects boundary samples equivalent to approximately 5% of the target dataset size. These samples represent attacks that overlap with benign ranges, such as slow port scanning, low-rate flooding, small DNS tunneling, and slow data exfiltration. Boundary samples are intended to reduce overfitting to extreme synthetic signatures.")

    doc.add_heading("3.3 Detection and Decision Fusion", level=2)
    add_para(doc,
             "The primary detector uses a RandomForest classifier trained on twelve flow features. During evaluation and live detection, deterministic C2 signatures are applied as post-processing overrides. The multi-agent version adds several analysis agents, including statistical, machine-learning, rule-signature, behavioral, temporal, protocol-specific, and flow-correlation analyzers. Their outputs are aggregated and passed through threshold management before alert generation.")

    doc.add_heading("4. Dataset Design", level=1)
    add_para(doc,
             "The prototype generates synthetic flow-level records with twelve model features, a binary label, and an attack_type field. The benign traffic set includes web browsing, DNS queries, file transfer, video streaming, and email. The current general generator covers fourteen anomaly classes.")
    add_table(doc,
              ["Class group", "Types"],
              [
                  ["Benign", "normal_web, normal_dns, normal_ftp, normal_stream, normal_email"],
                  ["Availability and DoS", "ddos, synflood, http_flood, slowloris, dns_amplification"],
                  ["Reconnaissance and credential attacks", "portscan, bruteforce, credential_stuffing"],
                  ["Command and control", "botnet_c2, dns_tunneling"],
                  ["Exfiltration and abuse", "exfiltration, cryptomining"],
                  ["Lateral/impact and local-network attacks", "ransomware, arp_spoofing"]
              ],
              [2500, 6860])
    add_para(doc,
             "A limitation is that the adaptive generator currently implements eleven attack classes, while the fixed generator has expanded to fourteen. This inconsistency should be resolved before final publication by extending adaptive weighting and boundary generation to cryptomining, DNS amplification, and credential stuffing.")

    doc.add_heading("5. Experimental Setup", level=1)
    add_para(doc,
             "The current prototype was evaluated using synthetic flow-feature data generated by the local pipeline. The training loop stops when F1 >= 0.92, recall >= 0.90, and precision >= 0.88. Test data are made more realistic by injecting boundary variants, high-load benign traffic, Gaussian feature noise, and limited label noise. The real-time phase uses simulated incoming CSV batches, where one batch corresponds to one generated flow file.")
    add_table(doc,
              ["Setting", "Value"],
              [
                  ["Primary model", "RandomForestClassifier"],
                  ["Input representation", "Twelve numerical flow features"],
                  ["Default training ratio", "65% benign, 35% anomalous"],
                  ["Evaluation targets", "F1 >= 0.92, recall >= 0.90, precision >= 0.88"],
                  ["Real-time simulation batch size", "50 flow records per incoming file"],
                  ["Signature override", "Known botnet C2 destination ports"],
                  ["Representative validation command", "run_pipeline.py --max-batches 1"]
              ],
              [2800, 6560])

    doc.add_heading("6. Preliminary Results", level=1)
    add_para(doc,
             "Table 4 summarizes representative results from the corrected prototype. These values should be interpreted as preliminary synthetic-flow results, not as evidence of production readiness on real network traffic.")
    add_table(doc,
              ["Run", "F1", "Recall", "Precision", "Phase-2 result"],
              [
                  ["Sequential pipeline, one-batch verification", "0.9702", "0.9607", "0.9799", "50 processed, 8 anomalies, 8 alerts"],
                  ["Direct best-model evaluation", "0.9674", "0.9525", "0.9827", "Evaluation script completed without SameFileError"],
                  ["Multi-agent pipeline, one-batch verification", "0.9718", "0.9583", "0.9857", "50 processed, 8 anomalies"]
              ],
              [3100, 1100, 1100, 1100, 2950])
    add_para(doc,
             "The sequential verification showed consistency between the detector session summary and alert summary: one 50-flow batch produced eight anomaly decisions and eight alert records. The botnet C2 signature override was also observed to produce high-confidence detections for known C2 ports in previous validation runs.")

    doc.add_heading("7. Discussion", level=1)
    add_para(doc,
             "The primary benefit of the framework is not the use of a particular classifier, but the operational coupling between evaluation feedback and future data generation. This coupling makes it possible to focus training effort on weak classes, rather than relying on a static and potentially misleading distribution. The multi-agent design further supports the integration of deterministic security evidence, such as known malicious ports, with probabilistic model output.")
    add_para(doc,
             "Several limitations remain. First, the system currently relies on synthetic flow-level data. Second, feature extraction from real packets or logs is not implemented as a production collector. Third, the file-based batch interface is suitable for experimentation but should be replaced by a queue-based microbatch architecture for high-rate traffic. Fourth, the attack taxonomy and adaptive generator are not yet fully synchronized after the recent expansion from eleven to fourteen anomaly classes.")

    doc.add_heading("8. Toward Real-World Deployment", level=1)
    add_para(doc,
             "To adapt this prototype for operational use, the file-based simulator should be replaced with a telemetry ingestion pipeline. A practical architecture would consist of packet or log collectors, a flow-feature aggregator, a message queue, detector workers, and asynchronous alert workers. The model should be deployed together with its scaler, feature schema, threshold configuration, signature rules, and model-version metadata.")
    add_numbered(doc, [
        "Implement a production feature extractor for Zeek, Suricata, NetFlow, firewall, or proxy logs.",
        "Replace CSV polling with Kafka, Redis Streams, NATS, or another queue mechanism.",
        "Run detection as microbatch inference workers rather than one process reading one file at a time.",
        "Move alert generation to asynchronous workers that publish to SIEM, webhook, Slack, email, or database sinks.",
        "Track latency, queue depth, dropped records, per-class alert counts, and model-version metadata.",
        "Validate against real benign traffic and controlled attack replay before making operational claims."
    ])

    doc.add_heading("9. Conclusion", level=1)
    add_para(doc,
             "This paper draft presented a feedback-driven adaptive MLOps framework for synthetic flow-level network anomaly detection. The prototype combines adaptive packet-flow generation, automated model training, per-attack evaluation, deterministic signature overrides, and multi-agent decision fusion. Preliminary results indicate that the framework can reach strong synthetic-data performance and maintain consistency between anomaly decisions and alert output. Future work should focus on real telemetry ingestion, attack-taxonomy synchronization, queue-based deployment, and evaluation on public and private network-security datasets.")

    doc.add_heading("References", level=1)
    refs = [
        "MITRE ATT&CK. Enterprise Matrix and Techniques. https://attack.mitre.org/",
        "MITRE ATT&CK. Network Denial of Service, T1498. https://attack.mitre.org/techniques/T1498/",
        "MITRE ATT&CK. Application Layer Protocol: DNS, T1071.004. https://attack.mitre.org/techniques/T1071/004/",
        "MITRE ATT&CK. Brute Force, T1110. https://attack.mitre.org/techniques/T1110/",
        "MITRE ATT&CK. Adversary-in-the-Middle, T1557. https://attack.mitre.org/techniques/T1557/",
        "To be completed: related work on network intrusion detection, data augmentation, synthetic traffic generation, and security MLOps."
    ]
    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(f"[{i}] ").bold = True
        p.add_run(ref)

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
